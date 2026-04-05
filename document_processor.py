"""
Módulo de Procesamiento de Documentos para Cic_IA
Extrae información de PDFs, imágenes, documentos Word, etc.
"""

import os
import logging
from typing import List, Dict, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Procesador de documentos para extracción de información"""
    
    @staticmethod
    def extract_from_pdf(file_path: str, max_pages: Optional[int] = None) -> Dict:
        """Extraer texto de un PDF"""
        try:
            import PyPDF2
            
            text_content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Obtener metadatos
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', 'Sin título'),
                        'author': pdf_reader.metadata.get('/Author', 'Desconocido'),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extraer texto
                pages_to_read = min(len(pdf_reader.pages), max_pages or len(pdf_reader.pages))
                
                for page_num in range(pages_to_read):
                    page = pdf_reader.pages[page_num]
                    text_content += f"\n--- Página {page_num + 1} ---\n"
                    text_content += page.extract_text()
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'pdf'
            }
        except ImportError:
            logger.warning("PyPDF2 no instalado, intentando con pdfplumber")
            return DocumentProcessor._extract_pdf_pdfplumber(file_path, max_pages)
        except Exception as e:
            logger.error(f"Error extrayendo PDF: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def _extract_pdf_pdfplumber(file_path: str, max_pages: Optional[int] = None) -> Dict:
        """Extraer PDF usando pdfplumber"""
        try:
            import pdfplumber
            
            text_content = ""
            metadata = {}
            
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    'pages': len(pdf.pages),
                    'title': 'Documento PDF'
                }
                
                pages_to_read = min(len(pdf.pages), max_pages or len(pdf.pages))
                
                for page_num in range(pages_to_read):
                    page = pdf.pages[page_num]
                    text_content += f"\n--- Página {page_num + 1} ---\n"
                    text_content += page.extract_text()
                    
                    # Extraer tablas si existen
                    tables = page.extract_tables()
                    if tables:
                        text_content += "\n[TABLAS DETECTADAS]\n"
                        for table in tables:
                            for row in table:
                                text_content += " | ".join(str(cell) for cell in row) + "\n"
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'pdf'
            }
        except ImportError:
            logger.error("pdfplumber no instalado")
            return {'success': False, 'error': 'PDF libraries not installed'}
        except Exception as e:
            logger.error(f"Error con pdfplumber: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict:
        """Extraer texto de un documento Word"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Extraer tablas
            if doc.tables:
                text_content += "\n[TABLAS DETECTADAS]\n"
                for table in doc.tables:
                    for row in table.rows:
                        text_content += " | ".join(cell.text for cell in row.cells) + "\n"
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'title': 'Documento Word'
            }
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'docx'
            }
        except ImportError:
            logger.error("python-docx no instalado")
            return {'success': False, 'error': 'python-docx not installed'}
        except Exception as e:
            logger.error(f"Error extrayendo DOCX: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def extract_from_image(file_path: str) -> Dict:
        """Extraer texto de una imagen usando OCR"""
        try:
            import pytesseract
            from PIL import Image
            
            # Abrir imagen
            image = Image.open(file_path)
            
            # Aplicar OCR
            text_content = pytesseract.image_to_string(image, lang='spa+eng')
            
            metadata = {
                'size': image.size,
                'format': image.format,
                'mode': image.mode
            }
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'image_ocr'
            }
        except ImportError:
            logger.warning("pytesseract no instalado, intentando con Pillow")
            return {'success': False, 'error': 'pytesseract not installed'}
        except Exception as e:
            logger.error(f"Error extrayendo imagen: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def extract_from_text(file_path: str) -> Dict:
        """Extraer contenido de un archivo de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            metadata = {
                'lines': len(text_content.split('\n')),
                'characters': len(text_content),
                'words': len(text_content.split())
            }
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'text'
            }
        except Exception as e:
            logger.error(f"Error extrayendo texto: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def extract_from_json(file_path: str) -> Dict:
        """Extraer contenido de un archivo JSON"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convertir a texto legible
            text_content = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                'type': 'json',
                'keys': list(data.keys()) if isinstance(data, dict) else 'array'
            }
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'json'
            }
        except Exception as e:
            logger.error(f"Error extrayendo JSON: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def extract_from_csv(file_path: str) -> Dict:
        """Extraer contenido de un archivo CSV"""
        try:
            import csv
            
            text_content = ""
            
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                
                for row in csv_reader:
                    text_content += " | ".join(row) + "\n"
            
            metadata = {
                'type': 'csv',
                'rows': len(text_content.split('\n'))
            }
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'source': 'csv'
            }
        except Exception as e:
            logger.error(f"Error extrayendo CSV: {e}")
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def extract_from_file(file_path: str) -> Dict:
        """Extraer contenido de cualquier archivo soportado"""
        file_ext = Path(file_path).suffix.lower()
        
        extractors = {
            '.pdf': DocumentProcessor.extract_from_pdf,
            '.docx': DocumentProcessor.extract_from_docx,
            '.doc': DocumentProcessor.extract_from_docx,
            '.txt': DocumentProcessor.extract_from_text,
            '.json': DocumentProcessor.extract_from_json,
            '.csv': DocumentProcessor.extract_from_csv,
            '.png': DocumentProcessor.extract_from_image,
            '.jpg': DocumentProcessor.extract_from_image,
            '.jpeg': DocumentProcessor.extract_from_image,
            '.gif': DocumentProcessor.extract_from_image,
        }
        
        extractor = extractors.get(file_ext)
        
        if not extractor:
            return {
                'success': False,
                'error': f'Formato no soportado: {file_ext}'
            }
        
        return extractor(file_path)


class ContentAnalyzer:
    """Analizador de contenido extraído"""
    
    @staticmethod
    def extract_summary(content: str, max_length: int = 500) -> str:
        """Extraer resumen del contenido"""
        try:
            import spacy
            nlp = spacy.load('es_core_news_sm')
            doc = nlp(content[:2000])  # Limitar a 2000 caracteres
            
            # Obtener oraciones
            sentences = list(doc.sents)
            
            if len(sentences) <= 3:
                return content[:max_length]
            
            # Calcular importancia
            sentence_scores = {}
            for sent in sentences:
                score = sum(token.vector_norm for token in sent if token.has_vector)
                sentence_scores[sent] = score
            
            # Seleccionar top 3 oraciones
            top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:3]
            top_sentences = sorted(top_sentences, key=lambda x: sentences.index(x[0]))
            
            summary = ' '.join([sent.text for sent, _ in top_sentences])
            return summary[:max_length]
        except:
            # Fallback: devolver primeras líneas
            lines = content.split('\n')
            return '\n'.join(lines[:3])[:max_length]
    
    @staticmethod
    def extract_entities(content: str) -> Dict[str, List[str]]:
        """Extraer entidades nombradas del contenido"""
        try:
            import spacy
            nlp = spacy.load('es_core_news_sm')
            doc = nlp(content[:2000])
            
            entities = {}
            for ent in doc.ents:
                if ent.label_ not in entities:
                    entities[ent.label_] = []
                if ent.text not in entities[ent.label_]:
                    entities[ent.label_].append(ent.text)
            
            return entities
        except:
            return {}
    
    @staticmethod
    def extract_keywords(content: str, num_keywords: int = 10) -> List[str]:
        """Extraer palabras clave del contenido"""
        try:
            from collections import Counter
            
            # Palabras comunes a ignorar
            stopwords = {
                'el', 'la', 'de', 'que', 'y', 'a', 'en', 'es', 'por', 'para',
                'con', 'una', 'un', 'los', 'las', 'del', 'al', 'este', 'ese',
                'esto', 'eso', 'aquello', 'más', 'menos', 'muy', 'bien', 'mal'
            }
            
            # Tokenizar
            words = content.lower().split()
            
            # Filtrar
            keywords = [
                w for w in words
                if len(w) > 3 and w not in stopwords and w.isalpha()
            ]
            
            # Contar frecuencias
            freq = Counter(keywords)
            
            return [kw for kw, _ in freq.most_common(num_keywords)]
        except:
            return []


# Ejemplo de uso
if __name__ == "__main__":
    # Procesar un PDF
    result = DocumentProcessor.extract_from_pdf("documento.pdf")
    
    if result['success']:
        print(f"Contenido extraído ({len(result['content'])} caracteres)")
        print(f"Metadatos: {result['metadata']}")
        
        # Analizar contenido
        summary = ContentAnalyzer.extract_summary(result['content'])
        keywords = ContentAnalyzer.extract_keywords(result['content'])
        
        print(f"Resumen: {summary}")
        print(f"Palabras clave: {keywords}")
    else:
        print(f"Error: {result['error']}")
