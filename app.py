"""
Bebé IA Pro Ultra - Sistema Completo
✅ Modelo estable (Gemma 2B) + Fallbacks
✅ Subida de documentos en conversación
✅ Aprendizaje continuo automático
✅ Autenticación con Google/GitHub
✅ Modos de respuesta: Rápido/Completo
✅ Persistencia de datos en PostgreSQL
"""
from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from authlib.integrations.flask_client import OAuth
import os
import json
import random
import re
import threading
import time
import hashlib
from datetime import datetime
from typing import List, Dict
import requests
import urllib.request
from werkzeug.utils import secure_filename
import PyPDF2
import docx

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'bebe-ia-secret-key-2024')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///bebe_ia.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

db = SQLAlchemy(app)
oauth = OAuth(app)

# Crear carpeta uploads
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ============ MODELOS DE BASE DE DATOS ============
class User(db.Model):
    """Usuarios registrados"""
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    name = db.Column(db.String(100))
    provider = db.Column(db.String(20))  # google, github, etc.
    provider_id = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Preferencias del usuario
    response_mode = db.Column(db.String(20), default='balanced')  # fast, balanced, complete
    interests = db.Column(db.Text, default='')  # Temas de interés separados por coma
    
    # Relaciones
    memories = db.relationship('Memory', backref='user', lazy=True)
    conversations = db.relationship('Conversation', backref='user', lazy=True)

class Memory(db.Model):
    """Memoria persistente de la IA"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    content = db.Column(db.Text, nullable=False)
    source = db.Column(db.String(50))  # wikipedia, arxiv, github, document, chat, etc.
    topic = db.Column(db.String(200))
    metadata_json = db.Column(db.Text)  # JSON adicional
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    access_count = db.Column(db.Integer, default=0)  # Cuántas veces se ha usado

class Conversation(db.Model):
    """Historial de conversaciones"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    user_message = db.Column(db.Text, nullable=False)
    bot_response = db.Column(db.Text, nullable=False)
    model_used = db.Column(db.String(100))
    response_mode = db.Column(db.String(20))
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

class Document(db.Model):
    """Documentos subidos por usuarios"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    filename = db.Column(db.String(255))
    original_name = db.Column(db.String(255))
    content_extracted = db.Column(db.Text)
    file_type = db.Column(db.String(10))  # pdf, docx, txt
    processed = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class LearningLog(db.Model):
    """Log de aprendizaje automático"""
    id = db.Column(db.Integer, primary_key=True)
    source = db.Column(db.String(50))  # wikipedia, arxiv, auto_search
    topic = db.Column(db.String(200))
    documents_found = db.Column(db.Integer)
    status = db.Column(db.String(20))  # success, error, pending
    error_message = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Crear tablas
with app.app_context():
    db.create_all()

# ============ CONFIGURACIÓN OAuth ============
# Google OAuth
app.config['GOOGLE_CLIENT_ID'] = os.environ.get('GOOGLE_CLIENT_ID', '')
app.config['GOOGLE_CLIENT_SECRET'] = os.environ.get('GOOGLE_CLIENT_SECRET', '')

if app.config['GOOGLE_CLIENT_ID']:
    oauth.register(
        name='google',
        client_id=app.config['GOOGLE_CLIENT_ID'],
        client_secret=app.config['GOOGLE_CLIENT_SECRET'],
        server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
        client_kwargs={'scope': 'openid email profile'}
    )

# GitHub OAuth
app.config['GITHUB_CLIENT_ID'] = os.environ.get('GITHUB_CLIENT_ID', '')
app.config['GITHUB_CLIENT_SECRET'] = os.environ.get('GITHUB_CLIENT_SECRET', '')

if app.config['GITHUB_CLIENT_ID']:
    oauth.register(
        name='github',
        client_id=app.config['GITHUB_CLIENT_ID'],
        client_secret=app.config['GITHUB_CLIENT_SECRET'],
        access_token_url='https://github.com/login/oauth/access_token',
        authorize_url='https://github.com/login/oauth/authorize',
        api_base_url='https://api.github.com/',
        client_kwargs={'scope': 'user:email'}
    )

# ============ CONFIGURACIÓN IA ============
class Config:
    # Modelos por orden de preferencia
    PRIMARY_MODEL = "google/gemma-2b-it"  # Rápido y estable
    SECONDARY_MODEL = "TinyLlama/TinyLlama-1.1B-Chat-v1.0"  # Ultra rápido
    TERTIARY_MODEL = "HuggingFaceH4/zephyr-7b-beta"  # Calidad si hay tiempo
    
    HF_API_TOKEN = os.environ.get('HF_API_TOKEN', '')
    
    # Aprendizaje automático
    AUTO_LEARN_INTERVAL = 900  # 15 minutos (más frecuente)
    MAX_DAILY_LEARNING = 30
    CONTINUOUS_SEARCH = True  # Buscar constantemente nuevos temas

# ============ CLIENTE API MULTI-MODELO ============
class HuggingFaceClient:
    """Cliente inteligente con múltiples modelos y fallbacks"""
    
    def __init__(self):
        self.headers = {
            "Authorization": f"Bearer {Config.HF_API_TOKEN}" if Config.HF_API_TOKEN else "",
            "Content-Type": "application/json"
        }
        self.current_model = Config.PRIMARY_MODEL
        self.model_status = {model: 'active' for model in [
            Config.PRIMARY_MODEL, 
            Config.SECONDARY_MODEL, 
            Config.TERTIARY_MODEL
        ]}
        print(f"🤖 Cliente API iniciado. Primario: {Config.PRIMARY_MODEL}")
    
    def generate(self, prompt: str, system_prompt: str = None, 
                 mode: str = 'balanced', max_retries: int = 3) -> tuple:
        """
        Generar respuesta con selección inteligente de modelo
        
        mode: 'fast' (rápido), 'balanced' (normal), 'complete' (detallado)
        """
        
        # Configurar parámetros según modo
        if mode == 'fast':
            max_tokens = 150
            temperature = 0.5
            models = [Config.SECONDARY_MODEL, Config.PRIMARY_MODEL]
        elif mode == 'complete':
            max_tokens = 1024
            temperature = 0.8
            models = [Config.TERTIARY_MODEL, Config.PRIMARY_MODEL]
        else:  # balanced
            max_tokens = 512
            temperature = 0.7
            models = [Config.PRIMARY_MODEL, Config.SECONDARY_MODEL]
        
        # Formatear prompt según modelo
        formatted_prompt = self._format_prompt(prompt, system_prompt, models[0])
        
        payload = {
            "inputs": formatted_prompt,
            "parameters": {
                "max_new_tokens": max_tokens,
                "temperature": temperature,
                "top_p": 0.9,
                "return_full_text": False
            }
        }
        
        # Intentar modelos en orden
        for model in models:
            for attempt in range(max_retries):
                try:
                    url = f"https://api-inference.huggingface.co/models/{model}"
                    response = requests.post(
                        url, 
                        headers=self.headers, 
                        json=payload, 
                        timeout=45 if mode != 'fast' else 20
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        if isinstance(result, list) and len(result) > 0:
                            text = result[0].get('generated_text', '')
                            text = self._clean_response(text, model)
                            if len(text) > 10:  # Respuesta válida
                                self.current_model = model
                                return text, model
                    
                    elif response.status_code == 503:
                        print(f"   ⏳ {model} cargando... intento {attempt+1}")
                        time.sleep(min(2 ** attempt, 10))  # Backoff exponencial
                        continue
                    
                    else:
                        print(f"   ❌ {model} error {response.status_code}")
                        
                except requests.Timeout:
                    print(f"   ⏱️ {model} timeout")
                    if mode == 'fast':
                        break  # No reintentar en modo rápido
                    time.sleep(2)
                    
                except Exception as e:
                    print(f"   ❌ {model} excepción: {e}")
            
            # Marcar modelo como problemático temporalmente
            self.model_status[model] = 'degraded'
        
        # Fallback final: respuesta local
        return self._local_fallback(prompt), "local_fallback"
    
    def _format_prompt(self, prompt: str, system: str, model: str) -> str:
        """Formatear según el modelo"""
        if 'gemma' in model.lower():
            if system:
                return f"<bos><start_of_turn>user\n{system}\n\n{prompt}<end_of_turn>\n<start_of_turn>model\n"
            return f"<bos><start_of_turn>user\n{prompt}<end_of_turn>\n<start_of_turn>model\n"
        
        elif 'tinyllama' in model.lower():
            if system:
                return f"### System:\n{system}\n\n### User:\n{prompt}\n\n### Assistant:\n"
            return f"### User:\n{prompt}\n\n### Assistant:\n"
        
        else:  # Mistral/Zephyr
            if system:
                return f"<s>[INST] {system}\n\n{prompt} [/INST]"
            return f"<s>[INST] {prompt} [/INST]"
    
    def _clean_response(self, text: str, model: str) -> str:
        """Limpiar respuesta de tokens especiales"""
        # Remover tokens de fin
        for token in ['<end_of_turn>', '</s>', '[/INST]', '### Assistant:', '### User:']:
            text = text.split(token)[0]
        
        # Limpiar espacios
        text = text.strip()
        
        # Remover repeticiones del prompt
        lines = text.split('\n')
        if len(lines) > 1:
            # Verificar si la primera línea es muy similar al prompt
            text = '\n'.join(lines[1:]) if len(lines[0]) < 50 else text
        
        return text.strip()
    
    def _local_fallback(self, prompt: str) -> str:
        """Respuesta local cuando todo falla"""
        prompt_lower = prompt.lower()
        
        # Respuestas contextuales básicas
        if any(w in prompt_lower for w in ['hola', 'hey', 'buenas']):
            return "¡Hola! 👋 Estoy funcionando en modo limitado, pero puedo conversar. ¿En qué te ayudo?"
        
        elif any(w in prompt_lower for w in ['gracias', 'gracias', 'ty']):
            return "¡De nada! 😊 Estoy aquí para ayudarte."
        
        elif any(w in prompt_lower for w in ['adiós', 'bye', 'hasta luego']):
            return "¡Hasta luego! 👋 Vuelve cuando quieras."
        
        elif '?' in prompt:
            return "Buena pregunta. Estoy procesándola... ¿Puedes ser más específico mientras tanto?"
        
        return "Entiendo. Cuéntame más sobre eso para poder ayudarte mejor."

# ============ PROCESADOR DE DOCUMENTOS ============
class DocumentProcessor:
    """Extraer texto de documentos subidos"""
    
    @staticmethod
    def process_pdf(filepath: str) -> str:
        """Extraer texto de PDF"""
        try:
            text = ""
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text[:10000]  # Limitar tamaño
        except Exception as e:
            return f"Error procesando PDF: {str(e)}"
    
    @staticmethod
    def process_docx(filepath: str) -> str:
        """Extraer texto de Word"""
        try:
            doc = docx.Document(filepath)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text[:10000]
        except Exception as e:
            return f"Error procesando DOCX: {str(e)}"
    
    @staticmethod
    def process_txt(filepath: str) -> str:
        """Leer archivo de texto"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                return file.read()[:10000]
        except Exception as e:
            return f"Error leyendo TXT: {str(e)}"
    
    @classmethod
    def process(cls, filepath: str, filename: str) -> str:
        """Procesar según extensión"""
        ext = filename.lower().split('.')[-1]
        
        processors = {
            'pdf': cls.process_pdf,
            'docx': cls.process_docx,
            'txt': cls.process_txt,
            'md': cls.process_txt,
        }
        
        processor = processors.get(ext, cls.process_txt)
        return processor(filepath)

# ============ RECOLECTORES DE CONOCIMIENTO ============
class WikipediaCollector:
    def search(self, query: str, lang: str = "es") -> List[Dict]:
        try:
            url = f"https://{lang}.wikipedia.org/w/api.php?action=query&list=search&srsearch={urllib.parse.quote(query)}&format=json&srlimit=3"
            with urllib.request.urlopen(url, timeout=10) as response:
                data = json.loads(response.read())
            
            results = []
            for item in data['query']['search']:
                content = self._get_content(item['title'], lang)
                if content:
                    results.append({
                        'title': item['title'],
                        'content': content[:2000],
                        'source': 'wikipedia',
                        'url': f"https://{lang}.wikipedia.org/wiki/{item['title'].replace(' ', '_')}"
                    })
            return results
        except Exception as e:
            print(f"Wiki error: {e}")
            return []
    
    def _get_content(self, title: str, lang: str) -> str:
        try:
            url = f"https://{lang}.wikipedia.org/w/api.php?action=query&prop=extracts&explaintext&exchars=3000&titles={urllib.parse.quote(title)}&format=json"
            with urllib.request.urlopen(url, timeout=10) as response:
                data = json.loads(response.read())
            pages = data['query']['pages']
            return list(pages.values())[0].get('extract', '')
        except:
            return ""

class ArxivCollector:
    def search(self, query: str, max_results: int = 3) -> List[Dict]:
        try:
            url = f"http://export.arxiv.org/api/query?search_query=all:{urllib.parse.quote(query)}&start=0&max_results={max_results}&sortBy=relevance"
            response = requests.get(url, timeout=15)
            
            entries = re.findall(r'<entry>(.*?)</entry>', response.text, re.DOTALL)
            results = []
            
            for entry in entries:
                title = re.search(r'<title>(.*?)</title>', entry, re.DOTALL)
                summary = re.search(r'<summary>(.*?)</summary>', entry, re.DOTALL)
                
                if title and summary:
                    results.append({
                        'title': re.sub(r'<[^>]+>', '', title.group(1)),
                        'content': re.sub(r'<[^>]+>', '', summary.group(1)),
                        'source': 'arxiv'
                    })
            return results
        except Exception as e:
            print(f"Arxiv error: {e}")
            return []

# ============ MOTOR DE APRENDIZAJE CONTINUO ============
class ContinuousLearningEngine:
    """Aprendizaje automático constante"""
    
    def __init__(self, hf_client: HuggingFaceClient):
        self.hf = hf_client
        self.wiki = WikipediaCollector()
        self.arxiv = ArxivCollector()
        self.running = True
        self.daily_count = 0
        self.last_reset = datetime.now()
        
        # Temas de interés general para aprender
        self.interest_topics = [
            "inteligencia artificial", "machine learning", "deep learning",
            "python programming", "neural networks", "data science",
            "natural language processing", "computer vision",
            "reinforcement learning", "transformers", "algorithms",
            "mathematics", "statistics", "physics", "chemistry",
            "biology", "history", "philosophy", "psychology",
            "economics", "medicine", "astronomy", "climate change"
        ]
        
        # Iniciar hilos de aprendizaje
        self.learning_thread = threading.Thread(target=self._continuous_learning, daemon=True)
        self.learning_thread.start()
        
        self.research_thread = threading.Thread(target=self._deep_research, daemon=True)
        self.research_thread.start()
        
        print("🔄 Aprendizaje continuo iniciado")
    
    def _continuous_learning(self):
        """Aprendizaje frecuente (cada 15 min)"""
        while self.running:
            # Reset diario
            if (datetime.now() - self.last_reset).days >= 1:
                self.daily_count = 0
                self.last_reset = datetime.now()
            
            if self.daily_count < Config.MAX_DAILY_LEARNING:
                topic = random.choice(self.interest_topics)
                self._learn_topic(topic)
                self.daily_count += 1
            
            time.sleep(Config.AUTO_LEARN_INTERVAL)
    
    def _deep_research(self):
        """Investigación profunda (cada hora)"""
        while self.running:
            # Buscar temas populares o tendencias
            trending = self._get_trending_topics()
            for topic in trending[:2]:  # 2 temas por hora
                self._deep_learn(topic)
            time.sleep(3600)  # 1 hora
    
    def _get_trending_topics(self) -> List[str]:
        """Obtener temas tendencia (simulado, podría ser Twitter/Reddit API)"""
        # En producción, podrías usar APIs de tendencias
        return random.sample(self.interest_topics, 3)
    
    def _learn_topic(self, topic: str):
        """Aprender un tema básico"""
        print(f"🎓 Aprendiendo: {topic}")
        
        # Wikipedia
        articles = self.wiki.search(topic)
        for article in articles[:2]:
            summary = self._summarize(article['content'])
            self._save_to_db(article['title'], summary, 'wikipedia', topic)
        
        # arXiv si es técnico
        if any(t in topic for t in ['ai', 'learning', 'network', 'data']):
            papers = self.arxiv.search(topic, max_results=2)
            for paper in papers:
                simplified = self._simplify_scientific(paper['content'])
                self._save_to_db(paper['title'], simplified, 'arxiv', topic)
        
        # Log
        log = LearningLog(
            source='auto_wiki_arxiv',
            topic=topic,
            documents_found=len(articles) + len(papers) if 'papers' in dir() else len(articles),
            status='success'
        )
        db.session.add(log)
        db.session.commit()
    
    def _deep_learn(self, topic: str):
        """Aprendizaje profundo con análisis"""
        print(f"🔬 Investigando a fondo: {topic}")
        
        # Buscar más contenido
        articles = self.wiki.search(topic, lang='en')  # También en inglés
        articles.extend(self.wiki.search(topic, lang='es'))
        
        # Analizar y crear síntesis
        all_content = "\n\n".join([a['content'] for a in articles[:3]])
        synthesis = self._create_synthesis(topic, all_content)
        
        self._save_to_db(
            f"Síntesis: {topic}",
            synthesis,
            'deep_research',
            topic
        )
    
    def _summarize(self, text: str) -> str:
        """Resumir con modelo"""
        prompt = f"Resume esto en 3 oraciones:\n\n{text[:1500]}\n\nResumen:"
        result, _ = self.hf.generate(prompt, mode='fast')
        return result[:500]
    
    def _simplify_scientific(self, text: str) -> str:
        """Simplificar texto científico"""
        prompt = f"Explica esto simplemente:\n\n{text[:1000]}\n\nExplicación:"
        result, _ = self.hf.generate(prompt, mode='balanced')
        return result[:600]
    
    def _create_synthesis(self, topic: str, content: str) -> str:
        """Crear síntesis de múltiples fuentes"""
        prompt = f"Crea un resumen completo sobre '{topic}' basado en:\n\n{content[:2000]}\n\nSíntesis:"
        result, _ = self.hf.generate(prompt, mode='complete')
        return result[:1000]
    
    def _save_to_db(self, title: str, content: str, source: str, topic: str):
        """Guardar en base de datos"""
        memory = Memory(
            content=f"{title}\n\n{content}",
            source=source,
            topic=topic,
            metadata_json=json.dumps({
                'title': title,
                'learned_at': datetime.now().isoformat()
            })
        )
        db.session.add(memory)
        db.session.commit()
        print(f"   ✅ Guardado: {title[:50]}...")
    
    def force_learn(self, source: str, query: str, user_id: int = None) -> str:
        """Aprendizaje forzado por usuario"""
        try:
            if source == 'wikipedia':
                articles = self.wiki.search(query)
                for article in articles[:3]:
                    self._save_to_db(article['title'], article['content'], 'wikipedia', query)
                return f"✅ Aprendí {len(articles)} artículos de Wikipedia sobre '{query}'"
            
            elif source == 'arxiv':
                papers = self.arxiv.search(query, max_results=3)
                for paper in papers:
                    self._save_to_db(paper['title'], paper['content'], 'arxiv', query)
                return f"✅ Aprendí {len(papers)} papers de arXiv sobre '{query}'"
            
            elif source == 'research':
                self._deep_learn(query)
                return f"🔬 Investigación profunda completada sobre '{query}'"
            
            return "❌ Fuente no válida. Usa: wikipedia, arxiv, research"
            
        except Exception as e:
            return f"❌ Error: {str(e)}"

# ============ BEbÉ IA COMPLETA ============
class BebeIAUltra:
    """Sistema completo de IA"""
    
    def __init__(self):
        print("=" * 70)
        print("🚀 BEbÉ IA PRO ULTRA")
        print("Multi-modelo + Auth + Documentos + Aprendizaje Continuo")
        print("=" * 70)
        
        self.hf = HuggingFaceClient()
        self.learner = ContinuousLearningEngine(self.hf)
        
        print("\n✅ Sistema iniciado")
        print(f"   📚 {Memory.query.count()} memorias en base de datos")
        print("   🔄 Aprendizaje continuo: ACTIVO")
    
    def chat(self, user_input: str, user_id: int = None, 
             response_mode: str = 'balanced') -> Dict:
        """Procesar mensaje con contexto personalizado"""
        
        # Obtener preferencias del usuario
        user = User.query.get(user_id) if user_id else None
        
        # Buscar memorias relevantes (globales + del usuario)
        memories = self._search_memories(user_input, user_id)
        context = self._build_context(memories, user)
        
        # Detectar si hay documentos pendientes del usuario
        pending_docs = Document.query.filter_by(
            user_id=user_id, 
            processed=False
        ).all()
        
        doc_context = ""
        for doc in pending_docs:
            doc_context += f"\n[Documento: {doc.original_name}]\n{doc.content_extracted[:500]}\n"
            doc.processed = True  # Marcar como usado
        db.session.commit()
        
        # Construir prompt completo
        system = self._build_system_prompt(user)
        
        full_prompt = f"""{context}

{doc_context}

Pregunta del usuario: {user_input}

Responde de manera útil y natural:"""
        
        # Generar respuesta
        response, model_used = self.hf.generate(
            full_prompt, 
            system, 
            mode=response_mode
        )
        
        # Guardar conversación
        conv = Conversation(
            user_id=user_id,
            user_message=user_input,
            bot_response=response,
            model_used=model_used,
            response_mode=response_mode
        )
        db.session.add(conv)
        
        # Guardar en memoria para futuro
        self._save_exchange(user_input, response, user_id)
        
        db.session.commit()
        
        return {
            'response': response,
            'model_used': model_used,
            'mode': response_mode,
            'sources_used': list(set([m.source for m in memories])),
            'memories_found': len(memories),
            'documents_processed': len(pending_docs),
            'total_memories': Memory.query.count()
        }
    
    def _search_memories(self, query: str, user_id: int = None) -> List[Memory]:
        """Búsqueda inteligente en memoria"""
        # Palabras clave de la consulta
        keywords = set(re.findall(r'\b\w+\b', query.lower()))
        
        # Buscar en todas las memorias
        all_memories = Memory.query.all()
        
        scored = []
        for mem in all_memories:
            score = 0
            mem_words = set(re.findall(r'\b\w+\b', mem.content.lower()))
            
            # Coincidencia de palabras
            matches = len(keywords & mem_words)
            score += matches * 2
            
            # Bonus si es del mismo usuario
            if user_id and mem.user_id == user_id:
                score += 5
            
            # Bonus por recencia
            days_old = (datetime.now() - mem.created_at).days
            score += max(0, 10 - days_old)
            
            # Bonus por frecuencia de uso
            score += min(mem.access_count, 5)
            
            if score > 2:
                scored.append((score, mem))
                mem.access_count += 1
        
        scored.sort(reverse=True, key=lambda x: x[0])
        return [m for _, m in scored[:5]]
    
    def _build_context(self, memories: List[Memory], user: User) -> str:
        """Construir contexto de conversación"""
        if not memories:
            return "Contexto: No hay información previa específica."
        
        parts = ["Información relevante de mi base de conocimiento:"]
        for mem in memories:
            source_icon = {
                'wikipedia': '📚',
                'arxiv': '📄',
                'github': '💻',
                'document': '📄',
                'deep_research': '🔬',
                'chat': '💬'
            }.get(mem.source, '📄')
            
            parts.append(f"{source_icon} [{mem.source}] {mem.content[:300]}...")
        
        # Agregar intereses del usuario si existen
        if user and user.interests:
            parts.append(f"\nIntereses del usuario: {user.interests}")
        
        return "\n\n".join(parts)
    
    def _build_system_prompt(self, user: User) -> str:
        """Construir prompt de sistema personalizado"""
        base = """Eres Bebé IA Pro Ultra, un asistente inteligente que:
- Aprende continuamente de Wikipedia, papers científicos y documentos
- Recuerda conversaciones previas y preferencias de usuarios
- Adapta sus respuestas según el modo solicitado (rápido/balanceado/completo)
- Puede procesar documentos subidos por el usuario"""
        
        if user:
            base += f"\n- Estás conversando con: {user.name or user.email}"
            if user.interests:
                base += f"\n- Sus intereses incluyen: {user.interests}"
        
        return base
    
    def _save_exchange(self, user_msg: str, bot_msg: str, user_id: int):
        """Guardar intercambio en memoria"""
        memory = Memory(
            user_id=user_id,
            content=f"Q: {user_msg}\nA: {bot_msg}",
            source='chat',
            topic=self._extract_topic(user_msg)
        )
        db.session.add(memory)
    
    def _extract_topic(self, text: str) -> str:
        """Extraer tema principal del texto"""
        # Simple: primeras palabras significativas
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        return ' '.join(words[:3]) if words else 'general'
    
    def process_document(self, file_path: str, original_name: str, 
                        user_id: int = None) -> str:
        """Procesar documento subido"""
        # Extraer texto
        content = DocumentProcessor.process(file_path, original_name)
        
        # Guardar en DB
        doc = Document(
            user_id=user_id,
            filename=file_path,
            original_name=original_name,
            content_extracted=content,
            file_type=original_name.split('.')[-1].lower(),
            processed=False  # Se procesará en la siguiente conversación
        )
        db.session.add(doc)
        
        # También guardar en memoria general
        summary = self._summarize_document(content)
        memory = Memory(
            user_id=user_id,
            content=f"Documento: {original_name}\n\nResumen: {summary}\n\nContenido: {content[:2000]}",
            source='document',
            topic=original_name,
            metadata_json=json.dumps({
                'filename': original_name,
                'size': len(content)
            })
        )
        db.session.add(memory)
        db.session.commit()
        
        return f"✅ Documento '{original_name}' procesado. Lo usaré en nuestra próxima conversación."
    
    def _summarize_document(self, content: str) -> str:
        """Crear resumen del documento"""
        prompt = f"Resume este documento en 3 oraciones:\n\n{content[:2000]}\n\nResumen:"
        result, _ = self.hf.generate(prompt, mode='fast')
        return result[:400]

# ============ INICIALIZACIÓN ============
bebe = BebeIAUltra()

# ============ RUTAS DE AUTENTICACIÓN ============
@app.route('/login')
def login():
    """Página de login con opciones"""
    return '''
    <h1>Bebé IA Pro - Login</h1>
    <p>Elige cómo conectarte:</p>
    <a href="/login/google">🔵 Google</a><br><br>
    <a href="/login/github">⚫ GitHub</a><br><br>
    <a href="/">⏭️ Continuar sin login</a>
    '''

@app.route('/login/google')
def login_google():
    if 'google' not in oauth._clients:
        return "Google OAuth no configurado. Añade GOOGLE_CLIENT_ID y GOOGLE_CLIENT_SECRET."
    redirect_uri = url_for('auth_google', _external=True)
    return oauth.google.authorize_redirect(redirect_uri)

@app.route('/auth/google')
def auth_google():
    try:
        token = oauth.google.authorize_access_token()
        user_info = token.get('userinfo')
        
        # Buscar o crear usuario
        user = User.query.filter_by(email=user_info['email']).first()
        if not user:
            user = User(
                email=user_info['email'],
                name=user_info.get('name'),
                provider='google',
                provider_id=user_info['sub']
            )
            db.session.add(user)
            db.session.commit()
        
        session['user_id'] = user.id
        session['user_name'] = user.name or user.email
        return redirect('/')
    except Exception as e:
        return f"Error de autenticación: {str(e)}"

@app.route('/login/github')
def login_github():
    if 'github' not in oauth._clients:
        return "GitHub OAuth no configurado. Añade GITHUB_CLIENT_ID y GITHUB_CLIENT_SECRET."
    redirect_uri = url_for('auth_github', _external=True)
    return oauth.github.authorize_redirect(redirect_uri)

@app.route('/auth/github')
def auth_github():
    try:
        token = oauth.github.authorize_access_token()
        resp = oauth.github.get('user', token=token)
        profile = resp.json()
        
        email = profile.get('email') or f"github_{profile['id']}@example.com"
        
        user = User.query.filter_by(provider_id=str(profile['id'])).first()
        if not user:
            user = User(
                email=email,
                name=profile.get('name') or profile.get('login'),
                provider='github',
                provider_id=str(profile['id'])
            )
            db.session.add(user)
            db.session.commit()
        
        session['user_id'] = user.id
        session['user_name'] = user.name or user.email
        return redirect('/')
    except Exception as e:
        return f"Error de autenticación: {str(e)}"

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/')

@app.route('/profile')
def profile():
    """Ver/editar perfil del usuario"""
    if 'user_id' not in session:
        return redirect('/login')
    
    user = User.query.get(session['user_id'])
    if not user:
        session.clear()
        return redirect('/login')
    
    return jsonify({
        'email': user.email,
        'name': user.name,
        'response_mode': user.response_mode,
        'interests': user.interests,
        'memories_count': len(user.memories),
        'conversations_count': len(user.conversations)
    })

@app.route('/profile/update', methods=['POST'])
def update_profile():
    """Actualizar preferencias del usuario"""
    if 'user_id' not in session:
        return jsonify({'error': 'No autenticado'}), 401
    
    data = request.json
    user = User.query.get(session['user_id'])
    
    if 'response_mode' in data:
        user.response_mode = data['response_mode']  # fast, balanced, complete
    if 'interests' in data:
        user.interests = data['interests']
    if 'name' in data:
        user.name = data['name']
    
    db.session.commit()
    return jsonify({'message': 'Perfil actualizado'})

# ============ RUTAS PRINCIPALES ============
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/chat', methods=['POST'])
def chat():
    data = request.json
    user_id = session.get('user_id')
    
    # Obtener modo de respuesta (del usuario o del request)
    user = User.query.get(user_id) if user_id else None
    mode = data.get('mode', user.response_mode if user else 'balanced')
    
    result = bebe.chat(
        data.get('message', ''),
        user_id=user_id,
        response_mode=mode
    )
    
    return jsonify(result)

@app.route('/upload', methods=['POST'])
def upload_document():
    """Subir documento para que la IA aprenda"""
    if 'document' not in request.files:
        return jsonify({'error': 'No se envió archivo'}), 400
    
    file = request.files['document']
    if file.filename == '':
        return jsonify({'error': 'Nombre vacío'}), 400
    
    user_id = session.get('user_id')
    
    # Guardar archivo
    filename = secure_filename(f"{datetime.now().timestamp()}_{file.filename}")
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    # Procesar
    result = bebe.process_document(filepath, file.filename, user_id)
    
    return jsonify({'message': result})

@app.route('/learn', methods=['POST'])
def force_learn():
    """Forzar aprendizaje manual"""
    data = request.json
    user_id = session.get('user_id')
    
    result = bebe.learner.force_learn(
        data.get('source', 'wikipedia'),
        data.get('query', ''),
        user_id
    )
    return jsonify({'message': result})

@app.route('/status', methods=['GET'])
def status():
    """Estado del sistema"""
    user_id = session.get('user_id')
    user = User.query.get(user_id) if user_id else None
    
    return jsonify({
        'authenticated': user_id is not None,
        'user': {
            'name': session.get('user_name'),
            'email': user.email if user else None,
            'mode': user.response_mode if user else 'balanced'
        } if user else None,
        'total_memories': Memory.query.count(),
        'total_conversations': Conversation.query.count(),
        'today_learned': bebe.learner.daily_count,
        'learning_active': bebe.learner.running,
        'models_available': list(bebe.hf.model_status.keys())
    })

@app.route('/memories', methods=['GET'])
def get_memories():
    """Ver memorias del usuario o globales"""
    user_id = session.get('user_id')
    
    query = Memory.query
    if user_id and request.args.get('mine') == 'true':
        query = query.filter_by(user_id=user_id)
    
    memories = query.order_by(Memory.created_at.desc()).limit(20).all()
    
    return jsonify([{
        'id': m.id,
        'source': m.source,
        'topic': m.topic,
        'content': m.content[:500] + '...' if len(m.content) > 500 else m.content,
        'created_at': m.created_at.isoformat(),
        'access_count': m.access_count
    } for m in memories])

@app.route('/history', methods=['GET'])
def get_history():
    """Historial de conversaciones del usuario"""
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'error': 'Autenticación requerida'}), 401
    
    conversations = Conversation.query.filter_by(user_id=user_id)\
        .order_by(Conversation.timestamp.desc())\
        .limit(50).all()
    
    return jsonify([{
        'user': c.user_message,
        'bot': c.bot_response,
        'model': c.model_used,
        'mode': c.response_mode,
        'timestamp': c.timestamp.isoformat()
    } for c in conversations])

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
